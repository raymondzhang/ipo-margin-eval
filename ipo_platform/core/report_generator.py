# -*- coding: utf-8 -*-
"""
IPO评分报告生成器

支持格式:
- DOCX (Word文档)
- HTML (网页)

报告内容:
- IPO基本信息
- 7维度评分拆解（含雷达图）
- 投资建议与风险提示
- 基石投资者明细
- 财务数据摘要
"""

import json
import base64
import io
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

import sys
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from ipo_platform.models.database import get_db

# 可选依赖
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from jinja2 import Template
    JINJA_AVAILABLE = True
except ImportError:
    JINJA_AVAILABLE = False

try:
    import plotly.graph_objects as go
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False


class IPOReportGenerator:
    """IPO评分报告生成器"""
    
    DIM_NAMES = {
        'profitability': '盈利能力',
        'allocation': '配售结构',
        'cornerstone': '基石投资',
        'pricing': '估值定价',
        'stabilization': '稳价机制',
        'q1_break_rate': 'Q1破发率',
        'hsi_monthly': 'HSI月度涨跌',
    }
    
    def __init__(self):
        self.db = get_db()
    
    def load_data(self, stock_code: str) -> Dict:
        """加载IPO完整数据"""
        base = self.db.fetchone("SELECT * FROM ipo_base WHERE stock_code = ?", (stock_code,))
        score = self.db.fetchone(
            "SELECT * FROM ipo_scores WHERE stock_code = ? ORDER BY scored_at DESC LIMIT 1",
            (stock_code,)
        )
        financials = self.db.fetchone("SELECT * FROM ipo_financials WHERE stock_code = ?", (stock_code,))
        cornerstones = self.db.fetchall("SELECT * FROM ipo_cornerstones WHERE stock_code = ?", (stock_code,))
        
        base = dict(base) if base else {}
        score = dict(score) if score else {}
        financials = dict(financials) if financials else {}
        cornerstones = [dict(c) for c in cornerstones]
        
        # 解析evidence_json
        dimensions = {}
        if score and score.get('evidence_json'):
            try:
                evidence = json.loads(score['evidence_json'])
                dimensions = evidence.get('dimensions', {})
            except:
                pass
        
        return {
            'base': base,
            'score': score,
            'financials': financials,
            'cornerstones': cornerstones,
            'dimensions': dimensions,
        }
    
    def _generate_radar_base64(self, dimensions: Dict) -> Optional[str]:
        """生成评分雷达图并返回base64编码"""
        if not PLOTLY_AVAILABLE or not dimensions:
            return None
        
        categories = list(dimensions.keys())
        scores = [dimensions[c]['score'] for c in categories]
        max_scores = [dimensions[c]['max'] for c in categories]
        
        labels = [self.DIM_NAMES.get(c, c) for c in categories]
        
        fig = go.Figure()
        fig.add_trace(go.Scatterpolar(
            r=scores + [scores[0]],
            theta=labels + [labels[0]],
            fill='toself',
            name='得分',
            line_color='#1f77b4',
            fillcolor='rgba(31, 119, 180, 0.3)'
        ))
        fig.add_trace(go.Scatterpolar(
            r=max_scores + [max_scores[0]],
            theta=labels + [labels[0]],
            name='满分',
            line_color='rgba(200,200,200,0.5)',
            line_dash='dash',
            fill=None
        ))
        fig.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, max(max_scores)])),
            showlegend=True,
            height=400,
            margin=dict(l=40, r=40, t=40, b=40)
        )
        
        img_bytes = fig.to_image(format="png", scale=2)
        return base64.b64encode(img_bytes).decode('utf-8')
    
    def generate_docx(self, stock_code: str, output_path: Optional[str] = None) -> str:
        """
        生成DOCX报告
        
        Args:
            stock_code: 股票代码
            output_path: 输出路径（默认 data/reports/{code}_report.docx）
        
        Returns:
            生成的文件路径
        """
        if not DOCX_AVAILABLE:
            raise ImportError("请安装 python-docx: python3 -m pip install python-docx")
        
        data = self.load_data(stock_code)
        base = data['base']
        score = data['score']
        financials = data['financials']
        cornerstones = data['cornerstones']
        dimensions = data['dimensions']
        
        if not base or not score:
            raise ValueError(f"数据库中未找到 {stock_code} 的完整数据")
        
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(10.5)
        
        # 标题
        title = doc.add_heading(f"IPO评分报告: {base.get('stock_name', '')} ({stock_code})", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生成日期
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
        
        doc.add_paragraph()
        
        # 基本信息表格
        doc.add_heading("一、基本信息", level=1)
        
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_items = [
            ("股票代码", stock_code),
            ("公司名称", base.get('stock_name', 'N/A')),
            ("行业", f"{base.get('industry', 'N/A')} ({base.get('industry_cat', 'N/A')})"),
            ("上市类型", base.get('reg_type', 'N/A')),
            ("保荐人", base.get('sponsor', 'N/A')),
            ("招股期", f"{base.get('ipo_start_date', 'N/A')} ~ {base.get('ipo_end_date', 'N/A')}"),
            ("上市日期", base.get('list_date', 'N/A')),
            ("招股价", f"{base.get('price_low', 'N/A')} ~ {base.get('price_high', 'N/A')} 港元"),
            ("募资额", f"{base.get('raise_amount', 'N/A')} 亿港元"),
            ("公开发售", f"{base.get('public_pct', 'N/A')}%"),
            ("国际配售", f"{base.get('intl_pct', 'N/A')}%"),
            ("稳价人", base.get('stabilizer', '无')),
            ("绿鞋", '有' if base.get('has_greenshoe') else '无'),
        ]
        
        for label, value in info_items:
            row = info_table.add_row().cells
            row[0].text = label
            row[1].text = str(value)
            row[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # 评分总览
        doc.add_heading("二、评分总览", level=1)
        
        total = score.get('total_score', 0)
        base_score = score.get('base_score', 0)
        category = score.get('category', 'unknown')
        cat_map = {'whitelist': '白名单', 'greylist': '灰名单', 'blacklist': '黑名单'}
        
        p = doc.add_paragraph()
        p.add_run(f"综合评分: ").bold = True
        run = p.add_run(f"{total}/110")
        run.bold = True
        run.font.size = Pt(16)
        if category == 'whitelist':
            run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
        elif category == 'greylist':
            run.font.color.rgb = RGBColor(0xFF, 0xC1, 0x07)
        else:
            run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)
        
        p = doc.add_paragraph()
        p.add_run(f"基础得分: {base_score}/100  |  分类: {cat_map.get(category, category)}")
        
        p = doc.add_paragraph()
        p.add_run(f"杠杆建议: ").bold = True
        p.add_run(score.get('leverage_advice', 'N/A'))
        
        p = doc.add_paragraph()
        p.add_run(f"风险提示: ").bold = True
        p.add_run(score.get('risk_warning', 'N/A'))
        
        doc.add_paragraph()
        
        # 维度明细
        doc.add_heading("三、评分维度拆解", level=1)
        
        dim_table = doc.add_table(rows=1, cols=4)
        dim_table.style = 'Light Grid Accent 1'
        hdr = dim_table.rows[0].cells
        hdr[0].text = '维度'
        hdr[1].text = '得分'
        hdr[2].text = '满分'
        hdr[3].text = '判定'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        
        for k, v in dimensions.items():
            row = dim_table.add_row().cells
            row[0].text = self.DIM_NAMES.get(k, k)
            row[1].text = str(v['score'])
            row[2].text = str(v['max'])
            row[3].text = v.get('desc', '')
        
        doc.add_paragraph()
        
        # 雷达图
        radar_b64 = self._generate_radar_base64(dimensions)
        if radar_b64:
            doc.add_heading("四、评分雷达图", level=1)
            img_stream = io.BytesIO(base64.b64decode(radar_b64))
            doc.add_picture(img_stream, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # 财务数据
        if financials:
            doc.add_heading("五、财务数据摘要", level=1)
            fin_table = doc.add_table(rows=0, cols=4)
            fin_table.style = 'Light Grid Accent 1'
            
            fin_rows = [
                ("营收 (百万元人民币)", 
                 financials.get('revenue_2023'), financials.get('revenue_2024'), financials.get('revenue_2025')),
                ("净利润 (百万元人民币)", 
                 financials.get('net_profit_2023'), financials.get('net_profit_2024'), financials.get('net_profit_2025')),
            ]
            
            hdr = fin_table.add_row().cells
            hdr[0].text = '指标'
            hdr[1].text = '2023'
            hdr[2].text = '2024'
            hdr[3].text = '2025'
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True
            
            for label, v1, v2, v3 in fin_rows:
                row = fin_table.add_row().cells
                row[0].text = label
                row[1].text = f"{v1:,.2f}" if v1 is not None else 'N/A'
                row[2].text = f"{v2:,.2f}" if v2 is not None else 'N/A'
                row[3].text = f"{v3:,.2f}" if v3 is not None else 'N/A'
            
            doc.add_paragraph()
        
        # 基石投资者
        if cornerstones:
            doc.add_heading("六、基石投资者", level=1)
            cs_table = doc.add_table(rows=1, cols=3)
            cs_table.style = 'Light Grid Accent 1'
            hdr = cs_table.rows[0].cells
            hdr[0].text = '投资者名称'
            hdr[1].text = '认购金额 (亿港元)'
            hdr[2].text = '知名机构'
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True
            
            for cs in cornerstones:
                row = cs_table.add_row().cells
                row[0].text = cs.get('investor_name', 'N/A')
                row[1].text = f"{cs.get('amount_hkd', 0):,.2f}" if cs.get('amount_hkd') else 'N/A'
                row[2].text = '是' if cs.get('is_star') else '否'
            
            doc.add_paragraph()
        
        # 免责声明
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("— 本报告仅供参考，不构成投资建议 —")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        
        # 保存
        if output_path is None:
            report_dir = Path(__file__).parent.parent / 'data' / 'reports'
            report_dir.mkdir(parents=True, exist_ok=True)
            output_path = report_dir / f"{stock_code.replace('.', '_')}_report.docx"
        
        doc.save(output_path)
        return str(output_path)
    
    def generate_html(self, stock_code: str, output_path: Optional[str] = None) -> str:
        """
        生成HTML报告
        
        Returns:
            生成的HTML文件路径
        """
        data = self.load_data(stock_code)
        base = data['base']
        score = data['score']
        financials = data['financials']
        cornerstones = data['cornerstones']
        dimensions = data['dimensions']
        
        if not base or not score:
            raise ValueError(f"数据库中未找到 {stock_code} 的完整数据")
        
        total = score.get('total_score', 0)
        base_score = score.get('base_score', 0)
        category = score.get('category', 'unknown')
        cat_map = {'whitelist': '白名单', 'greylist': '灰名单', 'blacklist': '黑名单'}
        cat_color = {'whitelist': '#28a745', 'greylist': '#ffc107', 'blacklist': '#dc3545'}
        
        radar_b64 = self._generate_radar_base64(dimensions)
        
        # 构建维度表格HTML
        dim_rows_html = ""
        for k, v in dimensions.items():
            dim_rows_html += f"""
            <tr>
                <td>{self.DIM_NAMES.get(k, k)}</td>
                <td>{v['score']}/{v['max']}</td>
                <td>{v.get('desc', '')}</td>
            </tr>"""
        
        # 财务表格HTML
        fin_html = ""
        if financials:
            fin_html = f"""
            <h2>五、财务数据摘要</h2>
            <table>
                <tr><th>指标</th><th>2023</th><th>2024</th><th>2025</th></tr>
                <tr>
                    <td>营收 (百万元人民币)</td>
                    <td>{financials.get('revenue_2023', 'N/A')}</td>
                    <td>{financials.get('revenue_2024', 'N/A')}</td>
                    <td>{financials.get('revenue_2025', 'N/A')}</td>
                </tr>
                <tr>
                    <td>净利润 (百万元人民币)</td>
                    <td>{financials.get('net_profit_2023', 'N/A')}</td>
                    <td>{financials.get('net_profit_2024', 'N/A')}</td>
                    <td>{financials.get('net_profit_2025', 'N/A')}</td>
                </tr>
            </table>"""
        
        # 基石表格HTML
        cs_html = ""
        if cornerstones:
            cs_rows = ""
            for cs in cornerstones:
                cs_rows += f"""
                <tr>
                    <td>{cs.get('investor_name', 'N/A')}</td>
                    <td>{cs.get('amount_hkd', 'N/A')}</td>
                    <td>{'是' if cs.get('is_star') else '否'}</td>
                </tr>"""
            cs_html = f"""
            <h2>六、基石投资者</h2>
            <table>
                <tr><th>投资者名称</th><th>认购金额 (亿港元)</th><th>知名机构</th></tr>
                {cs_rows}
            </table>"""
        
        # 雷达图HTML
        radar_html = ""
        if radar_b64:
            radar_html = f"""
            <h2>四、评分雷达图</h2>
            <div class="chart-container">
                <img src="data:image/png;base64,{radar_b64}" alt="评分雷达图">
            </div>"""
        
        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>IPO评分报告 - {base.get('stock_name', '')} ({stock_code})</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 40px 20px; background: #f5f5f5; }}
        .container {{ background: white; padding: 40px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ text-align: center; color: #1a1a1a; margin-bottom: 10px; }}
        .subtitle {{ text-align: center; color: #888; font-size: 0.9rem; margin-bottom: 30px; }}
        h2 {{ color: #1f77b4; border-bottom: 2px solid #1f77b4; padding-bottom: 8px; margin-top: 30px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th, td {{ padding: 10px 12px; text-align: left; border: 1px solid #ddd; }}
        th {{ background: #f8f9fa; font-weight: bold; }}
        tr:nth-child(even) {{ background: #fafafa; }}
        .score-box {{ text-align: center; padding: 25px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border-radius: 12px; margin: 20px 0; }}
        .score-number {{ font-size: 3rem; font-weight: bold; }}
        .score-label {{ font-size: 1rem; opacity: 0.9; }}
        .category-badge {{ display: inline-block; padding: 5px 20px; border-radius: 20px; font-weight: bold; color: white; background: {cat_color.get(category, '#666')}; }}
        .info-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin: 15px 0; }}
        .info-item {{ padding: 10px; background: #f8f9fa; border-radius: 6px; }}
        .info-label {{ font-weight: bold; color: #555; font-size: 0.9rem; }}
        .info-value {{ color: #1a1a1a; margin-top: 4px; }}
        .advice-box {{ padding: 15px; border-radius: 8px; margin: 10px 0; }}
        .advice-info {{ background: #e7f3ff; border-left: 4px solid #1f77b4; }}
        .advice-warn {{ background: #fff3cd; border-left: 4px solid #ffc107; }}
        .chart-container {{ text-align: center; margin: 20px 0; }}
        .chart-container img {{ max-width: 100%; height: auto; }}
        .footer {{ text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #eee; color: #999; font-size: 0.85rem; }}
    </style>
</head>
<body>
<div class="container">
    <h1>IPO评分报告</h1>
    <p class="subtitle">{base.get('stock_name', '')} ({stock_code}) | 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
    
    <div class="score-box">
        <div class="score-label">综合评分</div>
        <div class="score-number">{total}/110</div>
        <div style="margin-top: 10px;">
            <span style="opacity: 0.9;">基础分 {base_score}/100 | </span>
            <span class="category-badge">{cat_map.get(category, category)}</span>
        </div>
    </div>
    
    <h2>一、基本信息</h2>
    <div class="info-grid">
        <div class="info-item"><div class="info-label">股票代码</div><div class="info-value">{stock_code}</div></div>
        <div class="info-item"><div class="info-label">行业</div><div class="info-value">{base.get('industry', 'N/A')} ({base.get('industry_cat', 'N/A')})</div></div>
        <div class="info-item"><div class="info-label">上市类型</div><div class="info-value">{base.get('reg_type', 'N/A')}</div></div>
        <div class="info-item"><div class="info-label">保荐人</div><div class="info-value">{base.get('sponsor', 'N/A')}</div></div>
        <div class="info-item"><div class="info-label">招股期</div><div class="info-value">{base.get('ipo_start_date', 'N/A')} ~ {base.get('ipo_end_date', 'N/A')}</div></div>
        <div class="info-item"><div class="info-label">上市日期</div><div class="info-value">{base.get('list_date', 'N/A')}</div></div>
        <div class="info-item"><div class="info-label">招股价</div><div class="info-value">{base.get('price_low', 'N/A')} ~ {base.get('price_high', 'N/A')} 港元</div></div>
        <div class="info-item"><div class="info-label">募资额</div><div class="info-value">{base.get('raise_amount', 'N/A')} 亿港元</div></div>
        <div class="info-item"><div class="info-label">公开发售</div><div class="info-value">{base.get('public_pct', 'N/A')}%</div></div>
        <div class="info-item"><div class="info-label">国际配售</div><div class="info-value">{base.get('intl_pct', 'N/A')}%</div></div>
        <div class="info-item"><div class="info-label">稳价人</div><div class="info-value">{base.get('stabilizer', '无')}</div></div>
        <div class="info-item"><div class="info-label">绿鞋</div><div class="info-value">{'有' if base.get('has_greenshoe') else '无'}</div></div>
    </div>
    
    <h2>二、投资建议</h2>
    <div class="advice-box advice-info">
        <strong>杠杆建议:</strong> {score.get('leverage_advice', 'N/A')}
    </div>
    <div class="advice-box advice-warn">
        <strong>风险提示:</strong> {score.get('risk_warning', 'N/A')}
    </div>
    
    <h2>三、评分维度拆解</h2>
    <table>
        <tr><th>维度</th><th>得分</th><th>判定</th></tr>
        {dim_rows_html}
    </table>
    
    {radar_html}
    {fin_html}
    {cs_html}
    
    <div class="footer">
        — 本报告仅供参考，不构成投资建议 —
    </div>
</div>
</body>
</html>"""
        
        if output_path is None:
            report_dir = Path(__file__).parent.parent / 'data' / 'reports'
            report_dir.mkdir(parents=True, exist_ok=True)
            output_path = report_dir / f"{stock_code.replace('.', '_')}_report.html"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return str(output_path)


if __name__ == '__main__':
    gen = IPOReportGenerator()
    # 测试
    for code in ['07666.HK', '07630.HK']:
        try:
            docx_path = gen.generate_docx(code)
            html_path = gen.generate_html(code)
            print(f"{code}: DOCX={docx_path}, HTML={html_path}")
        except Exception as e:
            print(f"{code}: {e}")
