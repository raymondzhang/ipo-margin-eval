#!/usr/bin/env python3
"""
IPO Scoring Report Generator
Generates a formatted Word document (.docx) for HK IPO scoring reports.

Usage:
    python generate_report.py --input data.json --output report.docx

Or import functions and build programmatically:
    from generate_report import ReportBuilder
    builder = ReportBuilder()
    builder.add_title("IPO 标的质量评分报告")
    ...
    builder.save("report.docx")
"""

import json
import argparse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


class ReportBuilder:
    """Helper class to build IPO scoring reports with consistent formatting."""

    COLORS = {
        "heading": RGBColor(0x00, 0x00, 0x80),
        "green": RGBColor(0x00, 0x80, 0x00),
        "orange": RGBColor(0xFF, 0x66, 0x00),
        "red": RGBColor(0xCC, 0x00, 0x00),
        "blue": RGBColor(0x00, 0x66, 0xCC),
        "quote": RGBColor(0x40, 0x40, 0x40),
        "disclaimer": RGBColor(0x80, 0x80, 0x80),
    }

    def __init__(self):
        self.doc = Document()
        self._set_default_font()

    def _set_default_font(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _set_run_font(self, run, bold=False, italic=False, color=None, size=None):
        run.font.name = "Calibri"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size

    def add_main_title(self, text="IPO 标的质量评分报告"):
        title = self.doc.add_heading(text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            self._set_run_font(
                run, bold=True, size=Pt(22), color=self.COLORS["heading"]
            )
        self.doc.add_paragraph()

    def add_heading(self, text, level=1):
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            self._set_run_font(run)
            if level == 1:
                run.font.color.rgb = self.COLORS["heading"]
                run.font.size = Pt(18)
                run.bold = True
            elif level == 2:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                run.font.size = Pt(14)
                run.bold = True
        return heading

    def add_paragraph(self, text, bold=False, italic=False, color=None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_run_font(run, bold=bold, italic=italic, color=color)
        return p

    def add_bullet(self, text, bold=False):
        p = self.doc.add_paragraph(text, style="List Bullet")
        for run in p.runs:
            self._set_run_font(run, bold=bold)
        return p

    def add_quote(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        self._set_run_font(run, italic=True, color=self.COLORS["quote"])
        return p

    def add_score_display(self, score_text, color_key="green"):
        p = self.doc.add_paragraph()
        self._set_run_font(
            p.add_run(score_text),
            bold=True,
            size=Pt(14),
            color=self.COLORS[color_key],
        )
        return p

    def add_warning(self, text, color_key="orange"):
        p = self.doc.add_paragraph()
        self._set_run_font(p.add_run(text), color=self.COLORS[color_key])
        return p

    def add_basic_info(self, info: dict):
        """
        info keys: stock_code, company_cn, company_en, prospectus_date,
                   issue_price, offer_size, market_cap, industry
        """
        self.add_heading(
            f"标的：{info['company_cn']}（{info['company_en']}）", level=1
        )
        items = [
            f"股份代号：{info['stock_code']}",
            f"招股章程日期：{info['prospectus_date']}",
            f"发售价：每股H股{info['issue_price']}港元",
            f"全球发售规模：{info['offer_size']:,}股H股",
            f"市值：约{info['market_cap']}亿港元",
            f"行业：{info['industry']}",
        ]
        for item in items:
            p = self.doc.add_paragraph(item, style="List Bullet")
            for run in p.runs:
                self._set_run_font(run)

    def add_dimension_section(
        self,
        title: str,
        max_score: int,
        score: int,
        evidence_quotes: list,
        scoring_rules: list,
        reasoning: str,
        calculation: str = None,
        warning: str = None,
    ):
        """
        Add a complete scoring dimension section.
        """
        self.add_heading(f"{title}（满分{max_score}分）", level=1)

        self.add_heading("信息原文摘录", level=2)
        for quote in evidence_quotes:
            self.add_quote(quote)

        self.add_heading("评分规则对应依据", level=2)
        self.add_paragraph("评分标准：", bold=True)
        for rule in scoring_rules:
            self.add_bullet(rule)

        self.add_paragraph("判定：", bold=True)
        if calculation:
            self.add_paragraph(calculation)
        self.add_paragraph(reasoning, bold=True)

        self.add_heading("单项最终评分", level=2)
        color = "green" if score >= max_score * 0.8 else "orange" if score > 0 else "red"
        self.add_score_display(f"{score}分 / {max_score}分", color_key=color)
        if warning:
            self.add_warning(f"⚠️ {warning}", color_key=color)

    def add_summary_table(self, dimensions: list, total_max: int, total_score: int):
        """
        dimensions: list of dicts with keys: name, max, score
        """
        self.add_heading("总分汇总", level=1)
        table = self.doc.add_table(rows=len(dimensions) + 2, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        headers = ["评分维度", "满分", "得分", "得分率"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    self._set_run_font(run, bold=True)

        for i, dim in enumerate(dimensions, 1):
            row = table.rows[i]
            rate = f"{dim['score'] / dim['max']:.0%}"
            row.cells[0].text = dim["name"]
            row.cells[1].text = str(dim["max"])
            row.cells[2].text = str(dim["score"])
            row.cells[3].text = rate
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._set_run_font(run)

        total_row = table.rows[-1]
        total_row.cells[0].text = "合计"
        total_row.cells[1].text = str(total_max)
        total_row.cells[2].text = str(total_score)
        total_row.cells[3].text = f"{total_score / total_max:.1%}"
        for cell in total_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._set_run_font(run, bold=True)
        for paragraph in total_row.cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = self.COLORS["red"]
        for paragraph in total_row.cells[3].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = self.COLORS["red"]

    def add_classification_section(
        self, whitelist: bool, greylist: bool, blacklist_conditions: list, notes: str
    ):
        self.add_heading("IPO分类评估（白名单 / 灰名单 / 黑名单）", level=1)
        self.add_heading("分类标准", level=2)

        class_table = self.doc.add_table(rows=4, cols=2)
        class_table.style = "Light Shading Accent 1"
        ch = class_table.rows[0].cells
        ch[0].text = "分类"
        ch[1].text = "判定标准"
        for cell in ch:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._set_run_font(run, bold=True)

        data = [
            ["白名单", "得分≥80分（基础分），可直接开放打新融资额度"],
            ["灰名单", "18A与18C上市企业"],
            [
                "黑名单",
                "(1) 基石投资人不超过3家或认购比例不超过20%；\n"
                "(2) 发行静态市盈率超过行业龙头20%以上；\n"
                "(3) 招股期间发生重大不利负面舆情",
            ],
        ]
        for row_data in data:
            row = class_table.add_row().cells
            row[0].text = row_data[0]
            row[1].text = row_data[1]
            for paragraph in row[0].paragraphs:
                for run in paragraph.runs:
                    self._set_run_font(run, bold=True)

        self.add_heading("分类判定", level=2)
        cond_names = [
            "基石投资人≤3家或认购比例≤20%",
            "发行静态市盈率超过行业龙头20%以上",
            "招股期间发生重大不利负面舆情",
        ]
        for i, (triggered, name) in enumerate(zip(blacklist_conditions, cond_names), 1):
            status = "已触发" if triggered else "未触发"
            self.add_paragraph(f"条件({i}) {name}：{status}")

        self.add_paragraph("", bold=True)
        p = self.doc.add_paragraph()
        self._set_run_font(
            p.add_run(notes), bold=True, color=self.COLORS["blue"]
        )

    def add_recommendation(
        self,
        fundamentals: list,
        structure: list,
        macro: list,
        risks: list,
        decision: dict,
        rating: str,
    ):
        """
        decision keys: open_margin, leverage, rationale, watch_points
        """
        self.add_heading("标的总结与孖展额度决策建议", level=1)

        self.add_heading("基本面评价", level=2)
        for item in fundamentals:
            self.add_bullet(item)

        self.add_heading("发行结构评价", level=2)
        for item in structure:
            self.add_bullet(item)

        self.add_heading("宏观环境评价", level=2)
        for item in macro:
            self.add_bullet(item)

        self.add_heading("风险因素提示", level=2)
        for item in risks:
            self.add_bullet(item)

        self.add_heading("孖展额度决策建议", level=2)
        dec_table = self.doc.add_table(rows=5, cols=2)
        dec_table.style = "Light Shading Accent 1"
        dh = dec_table.rows[0].cells
        dh[0].text = "建议维度"
        dh[1].text = "评估结论"
        for cell in dh:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._set_run_font(run, bold=True)

        dec_data = [
            ["是否开放孖展额度", decision["open_margin"]],
            ["建议孖展杠杆倍数", decision["leverage"]],
            ["理由", decision["rationale"]],
            ["重点关注时点", decision["watch_points"]],
        ]
        for row_data in dec_data:
            row = dec_table.add_row().cells
            row[0].text = row_data[0]
            row[1].text = row_data[1]

        self.add_heading("综合评级", level=2)
        self.add_paragraph(rating, bold=True, size=Pt(12), color=self.COLORS["blue"])

    def add_disclaimer(self, text: str):
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        self._set_run_font(
            p.add_run(text), italic=True, color=self.COLORS["disclaimer"], size=Pt(9)
        )

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self, path: str):
        self.doc.save(path)
        print(f"Report saved to: {path}")


def load_json_and_generate(json_path: str, output_path: str):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    builder = ReportBuilder()
    builder.add_main_title(data.get("title", "IPO 标的质量评分报告"))
    builder.add_basic_info(data["basic_info"])
    builder.add_page_break()

    for dim in data["dimensions"]:
        builder.add_dimension_section(
            title=dim["title"],
            max_score=dim["max_score"],
            score=dim["score"],
            evidence_quotes=dim.get("evidence_quotes", []),
            scoring_rules=dim.get("scoring_rules", []),
            reasoning=dim["reasoning"],
            calculation=dim.get("calculation"),
            warning=dim.get("warning"),
        )

    builder.add_summary_table(
        data["dimensions"],
        data.get("total_max", 110),
        data.get("total_score", 0),
    )

    builder.add_classification_section(
        whitelist=data["classification"]["whitelist"],
        greylist=data["classification"]["greylist"],
        blacklist_conditions=data["classification"]["blacklist_conditions"],
        notes=data["classification"]["notes"],
    )

    builder.add_recommendation(
        fundamentals=data["recommendation"]["fundamentals"],
        structure=data["recommendation"]["structure"],
        macro=data["recommendation"]["macro"],
        risks=data["recommendation"]["risks"],
        decision=data["recommendation"]["decision"],
        rating=data["recommendation"]["rating"],
    )

    builder.add_disclaimer(data.get("disclaimer", ""))
    builder.save(output_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate IPO Scoring Report")
    parser.add_argument("--input", required=True, help="Path to JSON data file")
    parser.add_argument("--output", required=True, help="Path to output .docx")
    args = parser.parse_args()
    load_json_and_generate(args.input, args.output)
