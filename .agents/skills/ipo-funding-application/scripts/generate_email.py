#!/usr/bin/env python3
"""
IPO Funding Application Email Generator
Generates internal funding application emails in Markdown and optionally Word.

Usage:
    python generate_email.py --input data.json --output email.md
    python generate_email.py --input data.json --output email.docx --format docx
"""

import json
import argparse
from pathlib import Path
from datetime import datetime


def load_json(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def generate_markdown(data: dict) -> str:
    """Generate Markdown email from template data."""
    d = data
    
    md = f"""# {d['stock_code']} {d['company_name']} — IPO孖展融資申請

> 生成日期：{datetime.now().strftime('%Y-%m-%d')}

---

## 第一部分：致 Celine 的審閱請求

Dear Celine,

我已與 {d.get('confirmed_with', 'Emily')} 確認，本次 **{d['stock_code']} {d['company_name']}** 我們最多可提供 **{d.get('max_funding_hkd', '港幣 50,000,000元')}** 作為IPO孖展融資總額，**資金成本為 {d.get('cost_rate', '2.27')}%**。在向管理層申請審批前，想先請你協助審閱以下草擬申請內容，如有任何意見或建議（特別是有關孖展利率、手續費或預批融資額度部分），請指示：

### 其他券商融資利率及手續費參考

| 券商 | 銀行融資利率 | 手續費 |
|:---|:---|:---|
"""
    
    brokers = d.get('broker_comparison', [
        {"name": "富途證券", "rate": "0", "fee": "HKD 100"},
        {"name": "華盛證券", "rate": "0", "fee": "HKD 100"},
    ])
    for b in brokers:
        md += f"| {b['name']} | {b['rate']} | {b['fee']} |\n"
    
    md += f"""
---

## 第二部分：致 之光總、Jerry 及 Frank 的批核申請

致之光總、Jerry及 Frank:

**{d['stock_code']} {d['company_name']}** 已通過香港交易所批准，已進入首次公開募股（IPO）階段。為配合本次IPO，本公司計劃向客戶提供新股認購貸款預約服務，現正式提交批核申請。申請詳情如下：

### 一、公司及市場簡介

{d.get('company_intro', '【請補充公司簡介】')}

### 二、孖展融資安排

{d['stock_code']} {d['company_name']} 首次公開招股，本公司將推出相關孖展融資服務，具體安排如下：

| 項目 | 安排 |
|:---|:---|
| **融資利率** | **{d.get('funding_rate', '0')}%** |
| 最低申請數量 | {d.get('min_units', '15')}股起 |
| 融資認購槓桿 | {d.get('leverage', '10')}倍（客戶需支付{d.get('client_ratio', '10')}%本金，公司提供{d.get('firm_ratio', '90')}%融資） |
| **融資認購手續費** | **{d.get('fee_hkd', 'HKD 100')}** |
| **預批融資額度** | **{d.get('approved_funding_hkd', '港幣 10,000,000元')}**（作為本次IPO的孖展融資總額） |

> ⚠️ **注意**：此處預批融資額度為 **{d.get('approved_funding_hkd', '港幣 10,000,000元')}**，與第一部分致 Celine 的 **{d.get('max_funding_hkd', '港幣 50,000,000元')}** 是否一致，請確認最終額度。

### 三、其他

有關該公司最新財務數據，可參考以下連結：

📎 [港交所公告 - {d.get('prospectus_date', '2026年4月20日')}]({d.get('hkex_url', 'https://www1.hkexnews.hk')})

"""
    
    # Optional scoring section
    if d.get('include_scoring') and d.get('scoring'):
        s = d['scoring']
        md += f"""
### 四、IPO標的質量評估（參考）

經內部評分模型評估，{d['company_name']} 的IPO質量評分如下：

| 評分維度 | 得分 |
|:---|:---|
| 過往3年盈利情況 | {s.get('profitability', '—')}/30 |
| 分配機制 | {s.get('allocation', '—')}/10 |
| 基石投資人 | {s.get('cornerstone', '—')}/15 |
| 定價情況（靜態PE） | {s.get('pricing', '—')}/20 |
| 穩價機制 | {s.get('stabilization', '—')}/10 |
| 市場宏觀—新股破發率 | {s.get('break_rate', '—')}/15 |
| 市場宏觀—恒生指數 | {s.get('hsi', '—')}/10 |
| **合計** | **{s.get('total', '—')}/110** |

IPO分類：**{s.get('classification', '—')}**
孖展建議：**{s.get('recommendation', '—')}**

"""
    
    md += """---

## 申請結語

懇請營運部依照上述內容完成系統配置，並批核本申請。

敬請審核批准。

---

*本申請由自動化系統生成，請人工核對關鍵數據後提交。*
"""
    
    return md


def generate_docx(data: dict, output_path: str):
    """Generate Word document email."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("Error: python-docx is required for DOCX output.")
        print("Install with: pip install python-docx")
        return False
    
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    
    d = data
    
    # Title
    title = doc.add_heading(f"{d['stock_code']} {d['company_name']} — IPO孖展融資申請", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()
    
    # Part 1
    doc.add_heading("第一部分：致 Celine 的審閱請求", level=1)
    doc.add_paragraph(f"Dear Celine,")
    p = doc.add_paragraph()
    p.add_run(f"我已與 {d.get('confirmed_with', 'Emily')} 確認，本次 ")
    p.add_run(f"{d['stock_code']} {d['company_name']}").bold = True
    p.add_run(f" 我們最多可提供 ")
    p.add_run(f"{d.get('max_funding_hkd', '港幣 50,000,000元')}").bold = True
    p.add_run(f" 作為IPO孖展融資總額，")
    p.add_run(f"資金成本為 {d.get('cost_rate', '2.27')}%").bold = True
    p.add_run("...")
    
    # Part 2
    doc.add_heading("第二部分：致 之光總、Jerry 及 Frank 的批核申請", level=1)
    doc.add_paragraph(f"致之光總、Jerry及 Frank:")
    
    doc.add_heading("一、公司及市場簡介", level=2)
    doc.add_paragraph(d.get('company_intro', '【請補充公司簡介】'))
    
    doc.add_heading("二、孖展融資安排", level=2)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = "Light Shading Accent 1"
    rows = [
        ("融資利率", f"{d.get('funding_rate', '0')}%"),
        ("最低申請數量", f"{d.get('min_units', '15')}股起"),
        ("融資認購槓桿", f"{d.get('leverage', '10')}倍"),
        ("融資認購手續費", d.get('fee_hkd', 'HKD 100')),
        ("預批融資額度", d.get('approved_funding_hkd', '港幣 10,000,000元')),
    ]
    for i, (key, val) in enumerate(rows, 1):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_heading("申請結語", level=2)
    doc.add_paragraph("懇請營運部依照上述內容完成系統配置，並批核本申請。")
    doc.add_paragraph("敬請審核批准。")
    
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")
    return True


def main():
    parser = argparse.ArgumentParser(description="Generate IPO Funding Application Email")
    parser.add_argument("--input", required=True, help="Path to JSON data file")
    parser.add_argument("--output", required=True, help="Output file path")
    parser.add_argument("--format", choices=["md", "docx", "both"], default="md", help="Output format")
    args = parser.parse_args()
    
    data = load_json(args.input)
    
    if args.format in ("md", "both"):
        md_path = args.output if args.output.endswith(".md") else args.output + ".md"
        md_content = generate_markdown(data)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"Markdown email saved to: {md_path}")
    
    if args.format in ("docx", "both"):
        docx_path = args.output if args.output.endswith(".docx") else args.output.replace(".md", ".docx") + ".docx"
        generate_docx(data, docx_path)


if __name__ == "__main__":
    main()
